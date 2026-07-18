"""
cr_engine.py — Moteur de génération des comptes rendus de DIAGNOSTIC (SDC).

Deux briques indépendantes :
  1. extraire()  : transcription note vocale + fiche chantier  ->  JSON structuré (Claude)
  2. generer()   : JSON structuré + template Word + photos      ->  .docx prêt à relire

Le template de l'agence (DIAG_MODEL_SDC.docx) n'est JAMAIS reconstruit :
il est ouvert, rempli, enregistré sous un autre nom. Styles, logo, cartouche
de pied de page et mise en page d'origine sont donc conservés à l'identique.
"""

from __future__ import annotations

import copy
import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

EXT_PHOTOS = {".jpg", ".jpeg", ".png", ".heic", ".webp"}
LARGEUR_PHOTO_CM = 8.3          # 2 photos par ligne sur A4 (marges ~0,6")
PHOTOS_PAR_LIGNE = 2

COULEUR_PRIORITE = {
    "URGENT": RGBColor(0xC0, 0x00, 0x00),
    "COURT TERME": RGBColor(0xBF, 0x62, 0x00),
    "MOYEN TERME": RGBColor(0x40, 0x40, 0x40),
}


# ---------------------------------------------------------------------------
# 1. EXTRACTION  (note vocale -> JSON)
# ---------------------------------------------------------------------------

def extraire(transcription: str, chantier: dict, photos: list[str],
             prompt_path: str = "prompt_diagnostic.md",
             model: str = "claude-sonnet-4-6") -> dict:
    """Appelle Claude pour transformer la note vocale en structure de CR.

    Nécessite la variable d'environnement ANTHROPIC_API_KEY.
    """
    import anthropic

    md = Path(prompt_path).read_text(encoding="utf-8")
    system = md.split("## SYSTEM", 1)[1].split("## USER", 1)[0].strip()

    presents = ", ".join(c["nom"] or c["role"] for c in chantier["contacts"] if c.get("present"))
    absents = ", ".join(c["nom"] or c["role"] for c in chantier["contacts"] if not c.get("present"))
    liste_photos = "\n".join(f"{i + 1}. {p}" for i, p in enumerate(photos)) or "(aucune)"

    user = f"""FICHE CHANTIER (source : Monday)
- SDC / immeuble : {chantier['sdc']}
- Adresse : {chantier['adresse']}
- Code immeuble : {chantier['code_immeuble']}
- N° d'affaire : {chantier['n_affaire']}
- MOA / syndic : {chantier['moa']}
- Date de la visite : {chantier['date_visite']}
- Présents : {presents}
- Absents : {absents}

FICHIERS PHOTOS (ordre chronologique)
{liste_photos}

TRANSCRIPTION DE LA NOTE VOCALE
\"\"\"
{transcription}
\"\"\""""

    client = anthropic.Anthropic()
    rep = client.messages.create(
        model=model,
        max_tokens=8000,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    txt = "".join(b.text for b in rep.content if b.type == "text").strip()
    txt = txt.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    return json.loads(txt)


# ---------------------------------------------------------------------------
# Helpers python-docx
# ---------------------------------------------------------------------------

def _paragraphes(doc: Document) -> list[Paragraph]:
    """Tous les paragraphes du corps, y compris ceux situés dans le content-control
    (page de garde + contacts + généralités sont dans un <w:sdt> dans ce template)."""
    return [Paragraph(p, doc) for p in doc.element.body.iter(qn("w:p"))]


def _trouver(doc: Document, debut: str) -> Paragraph:
    cible = debut.strip().lower()
    for p in _paragraphes(doc):
        if p.text.replace("\xa0", " ").strip().lower().startswith(cible):
            return p
    raise ValueError(f"Repère introuvable dans le template : {debut!r}")


def _completer(par: Paragraph, texte: str) -> None:
    """Ajoute du texte à la fin du paragraphe en reprenant le format du dernier run
    (utilisé pour « SDC … », « VISITE du … », « Code immeuble : … »)."""
    if not par.runs:
        par.add_run(texte)
        return
    src = par.runs[-1]
    run = par.add_run(texte)
    run.font.name = src.font.name
    run.font.size = src.font.size
    run.font.bold = src.font.bold
    run.font.color.rgb = src.font.color.rgb if src.font.color and src.font.color.rgb else None


def _inserer_apres(par: Paragraph, texte: str = "", *, style=None, gras=False,
                   italique=False, taille=10, couleur=None, espace_avant=0) -> Paragraph:
    new_p = copy.deepcopy(par._p)
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
            new_p.remove(child)
    par._p.addnext(new_p)
    np = Paragraph(new_p, par._parent)
    if style:
        np.style = style
    else:
        np.style = par.part.document.styles["Normal"]
    np.paragraph_format.space_before = Pt(espace_avant)
    np.paragraph_format.space_after = Pt(2)
    if texte:
        r = np.add_run(texte)
        r.font.bold = gras
        r.font.italic = italique
        r.font.size = Pt(taille)
        r.font.name = "Arial"
        if couleur is not None:
            r.font.color.rgb = couleur
    return np


def _ecrire_cellule(cell, texte: str, *, gras=False, taille=9, centre=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, ligne in enumerate(str(texte).split("\n")):
        if i:
            p = cell.add_paragraph()
            if centre:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ligne)
        r.font.bold = gras
        r.font.size = Pt(taille)
        r.font.name = "Arial"


# ---------------------------------------------------------------------------
# 2. GÉNÉRATION  (JSON -> .docx)
# ---------------------------------------------------------------------------

def generer(donnees: dict, chantier: dict, template: str, sortie: str,
            dossier_photos: str | None = None) -> str:
    doc = Document(template)

    # --- Page de garde --------------------------------------------------
    _completer(_trouver(doc, "SDC"), chantier["sdc"])
    _completer(_trouver(doc, "VISITE du"), _fr_date(chantier["date_visite"]))
    _completer(_trouver(doc, "Code immeuble"), str(chantier["code_immeuble"]))

    # --- Tableau des contacts -------------------------------------------
    _remplir_contacts(doc, chantier["contacts"])

    # --- GENERALITES ----------------------------------------------------
    ancre = _trouver(doc, "generalites")
    for para in reversed(donnees.get("generalites", [])):
        _inserer_apres(ancre, para, espace_avant=4)

    # --- OBSERVATIONS / RECOMMANDATIONS ---------------------------------
    ancre = _trouver(doc, "OBSERVATIONS")
    for obs in reversed(donnees.get("observations", [])):
        blocs: list[tuple] = []
        blocs.append((obs["titre"], dict(gras=True, taille=11, espace_avant=10)))
        prio = (obs.get("priorite") or "").upper()
        if prio:
            blocs.append((f"Priorité : {prio}",
                          dict(gras=True, taille=9, couleur=COULEUR_PRIORITE.get(prio))))
        if obs.get("constat"):
            blocs.append((f"Constat — {obs['constat']}", dict(taille=10)))
        if obs.get("analyse"):
            blocs.append((f"Analyse — {obs['analyse']}", dict(taille=10)))
        if obs.get("preconisation"):
            blocs.append((f"Préconisation — {obs['preconisation']}", dict(taille=10, gras=False)))
        if obs.get("photos_liees"):
            nums = ", ".join(f"n° {n}" for n in obs["photos_liees"])
            blocs.append((f"(cf. photos {nums})", dict(taille=9, italique=True)))
        for texte, kw in reversed(blocs):
            _inserer_apres(ancre, texte, **kw)

    # --- CONCLUSION ET RECOMMANDATIONS ÉNERGÉTIQUES ---------------------
    ancre = _trouver(doc, "Conclusion et recommandations")
    concl = donnees.get("conclusion", {})
    blocs = []
    for para in concl.get("synthese", []):
        blocs.append((para, dict(taille=10, espace_avant=4)))
    if concl.get("energetique"):
        blocs.append(("Volet énergétique", dict(gras=True, taille=10, espace_avant=8)))
        for para in concl["energetique"]:
            blocs.append((para, dict(taille=10)))
    for texte, kw in reversed(blocs):
        _inserer_apres(ancre, texte, **kw)

    # --- PHOTOS ----------------------------------------------------------
    if dossier_photos:
        _inserer_photos(doc, donnees.get("photos", []), dossier_photos)

    # --- Cartouche du pied de page ---------------------------------------
    _remplir_cartouche(doc, chantier)

    Path(sortie).parent.mkdir(parents=True, exist_ok=True)
    doc.save(sortie)
    return sortie


def _fr_date(iso: str) -> str:
    try:
        return datetime.fromisoformat(iso).strftime("%d/%m/%Y")
    except ValueError:
        return iso


def _tables(doc: Document) -> list[Table]:
    """Le tableau des contacts est situé dans le content-control : doc.tables ne le voit pas."""
    return [Table(t, doc) for t in doc.element.body.iter(qn("w:tbl"))]


def _remplir_contacts(doc: Document, contacts: list[dict]) -> None:
    tbl: Table = _tables(doc)[0]        # 6 col. : Organisme | Nom | Tél | Mail | Prés | Abs

    def ligne_modele(idx: int):
        return tbl.rows[idx]

    def remplir(row, c: dict) -> None:
        cells = row.cells
        _ecrire_cellule(cells[0], c.get("organisme", "") or c.get("role", ""), gras=True)
        _ecrire_cellule(cells[1], c.get("nom", ""))
        _ecrire_cellule(cells[2], c.get("telephone", ""))
        _ecrire_cellule(cells[3], c.get("email", ""))
        _ecrire_cellule(cells[4], "X" if c.get("present") else "", centre=True)
        _ecrire_cellule(cells[5], "" if c.get("present") else "X", centre=True)

    moa = [c for c in contacts if c.get("groupe", "MOA").upper() == "MOA"]
    moe = [c for c in contacts if c.get("groupe", "").upper() == "MOE"]

    # Le template fournit 2 lignes MOA (SYNDIC, CS) et 1 ligne MOE (GO ARCHITECTURE).
    lignes_moa = [tbl.rows[2], tbl.rows[3]]
    for i, c in enumerate(moa[:2]):
        remplir(lignes_moa[i], c)
    # Contacts MOA supplémentaires : on clone la dernière ligne MOA.
    ancre = tbl.rows[3]._tr
    for c in moa[2:]:
        tr = copy.deepcopy(tbl.rows[3]._tr)
        ancre.addnext(tr)
        ancre = tr
        remplir(tbl.rows[[r._tr for r in tbl.rows].index(tr)], c)

    for c in moe:
        remplir(tbl.rows[-1], c)


def _inserer_photos(doc: Document, photos_json: list[dict], dossier: str) -> None:
    dossier_p = Path(dossier)
    fichiers = [f for f in sorted(dossier_p.iterdir()) if f.suffix.lower() in EXT_PHOTOS]
    if not fichiers:
        return
    legendes = {p.get("fichier"): p.get("legende", "") for p in photos_json}

    ancre = _trouver(doc, "photos")
    nb_lignes = (len(fichiers) + PHOTOS_PAR_LIGNE - 1) // PHOTOS_PAR_LIGNE
    tbl = doc.add_table(rows=nb_lignes, cols=PHOTOS_PAR_LIGNE)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ancre._p.addnext(tbl._tbl)          # déplace le tableau juste sous le titre « photos »

    for i, fichier in enumerate(fichiers):
        cell = tbl.rows[i // PHOTOS_PAR_LIGNE].cells[i % PHOTOS_PAR_LIGNE]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        largeur = Cm(LARGEUR_PHOTO_CM)
        with Image.open(fichier) as im:
            if im.height > im.width:      # portrait : on limite la hauteur
                p.add_run().add_picture(str(fichier), height=Cm(LARGEUR_PHOTO_CM * 1.33))
            else:
                p.add_run().add_picture(str(fichier), width=largeur)
        legende = legendes.get(fichier.name, "")
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Photo n° {i + 1}" + (f" — {legende}" if legende else ""))
        r.font.size = Pt(8)
        r.font.italic = True
        r.font.name = "Arial"


def _remplir_cartouche(doc: Document, chantier: dict) -> None:
    footer = doc.sections[0].footer
    if not footer.tables:
        return
    t = footer.tables[0]
    _ecrire_cellule(t.rows[0].cells[5], _fr_date(chantier["date_visite"]), taille=8, centre=True)
    _ecrire_cellule(t.rows[1].cells[0], chantier["moa"], taille=8, centre=True)
    _ecrire_cellule(t.rows[1].cells[1], chantier["adresse"], taille=8, centre=True)
